# -*- encoding: utf-8 -*-
'''
@File    :   data_utils.py
@Time    :   2022/12/21 10:04:04
@Author  :   Zhifeng Li
@Contact :   li_zaaachary@outlook.com
@Desc    :   
'''
import os
import random
import json
import csv 
import logging
from collections import OrderedDict, defaultdict

try:
    import pandas as pd
    from docx import Document
except:
    print('warning: pandas or python-docx has not been installed.')


def load_data_txt(file_path):
    with open(file_path, 'r', encoding='utf-8')as f:
        lines = list(filter(lambda x:x, map(lambda x:x.strip(), f.readlines())))
    return lines


def load_data(file_path, mode='plain', df2list=False):
    '''
    load data from jsonl/tsf or plain
    mode: jsonl / tsf / plain
    '''
    result = []
    if mode == "feather":
        f = open(file_path, 'rb')
        feather_df = pd.read_feather(f)
        if not df2list:
            return feather_df
        else:
            for index, row in feather_df.iterrows():
                temp = []
                for key in row.keys():
                    temp.append(row[key])
                result.append(temp)

            return result
    elif mode == 'docx':
        doc = Document(file_path)
        paragraphs = doc.paragraphs
        result = []
        for par in paragraphs:
            result.append(par.text)
        return result

    f = open(file_path, 'r', encoding='utf-8')

    if mode == 'json':
        result =  json.load(f, object_pairs_hook=OrderedDict)
    elif mode == 'csv':
        reader = csv.reader(f)
        for item in reader:
            result.append(item)
    else:
        while True:
            line = f.readline()
            if not line:
                break
            if mode == 'jsonl':
                try:
                    line = json.loads(line, object_pairs_hook=OrderedDict)
                except:
                    print(line)
                    raise Exception('json load failed')
                result.append(line)
            elif mode == 'tsf':
                line = line.strip('\n').split('\t')
                result.append(line)
            elif mode == 'plain':
                line = line.strip()
                result.append(line)

    f.close()
    return result


def dump_data(target, file_path, mode='json', csv_row=None):
    if mode =='docx':
        document = Document()
        for line in target:
            if type(line) == str:
                document.add_paragraph(line)
            elif type(line) in [dict, OrderedDict, defaultdict]:
                for key, value in line.items():
                    document.add_paragraph(f"{key}: {value}")
                document.add_paragraph('\n')
            elif type(line) == list:
                for l in line:
                    document.add_paragraph(l)
            else:
                raise Exception('not implemented')
        if not file_path.endswith('.docx'):
            document.save(file_path + '.docx')
        else:
            document.save(file_path)
        return 
    elif mode == 'csv':
        assert csv_row is not None
        with open(file_path, 'w', encoding='utf-8') as csvfile:
            csvwriter = csv.writer(csvfile)
            csvwriter.writerow(csv_row)

            rows = []
            for line in target:
                row = []
                for key in csv_row:
                    row.append(line[key])
                rows.append(row)
            # import pdb; pdb.set_trace()
            csvwriter.writerows(rows)
        return

    f = open(file_path, 'w', encoding='utf-8')
    if mode == 'json':
        json.dump(target, f, ensure_ascii=False, indent=2)
    elif mode == 'tsf':
        for line in target:
            line = list(map(str, line))
            f.write('\t'.join(line) + '\n')
    elif mode == 'jsonl':
        for line in target:
            line = json.dumps(line, ensure_ascii=False)
            f.write(line+'\n')
    elif mode == 'plain':
        for line in target:
            line = str(line)
            f.write(line + '\n')

    f.close()


if __name__ == '__main__':
    print(load_data_txt('./dataspace/todo_keywords.txt'))